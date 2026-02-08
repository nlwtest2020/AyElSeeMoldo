"""
Market Research Parallel Analysis Engine.

Parses Word (.docx) and Excel (.xlsx) documents, extracts key themes,
and finds parallels across Moldova, Georgia, and Armenia markets.
"""

import re
from collections import Counter, defaultdict

import nltk
import pandas as pd
from docx import Document
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer

try:
    STOP_WORDS = set(stopwords.words("english"))
except LookupError:
    nltk.download("stopwords", quiet=True)
    STOP_WORDS = set(stopwords.words("english"))

LEMMATIZER = WordNetLemmatizer()

# Domain-specific keywords that signal market research themes
THEME_SEEDS = {
    "technology": ["it", "software", "tech", "digital", "programming", "developer",
                   "computer", "startup", "saas", "cloud", "ai", "data"],
    "language_education": ["english", "language", "course", "learning", "education",
                           "training", "certification", "skill", "esl", "tutor"],
    "employment": ["job", "employment", "career", "salary", "hiring", "workforce",
                   "unemployment", "labor", "remote", "freelance", "outsource"],
    "demographics": ["youth", "age", "population", "urban", "rural", "gender",
                     "millennial", "generation", "student", "graduate"],
    "economy": ["gdp", "growth", "economy", "market", "investment", "export",
                "import", "trade", "revenue", "income", "poverty"],
    "social_media": ["facebook", "instagram", "tiktok", "social", "media",
                     "online", "platform", "influencer", "content", "engagement"],
    "culture": ["culture", "tradition", "diaspora", "identity", "heritage",
                "community", "family", "values", "religion"],
}

MARKETS = ["moldova", "georgia", "armenia"]


def extract_text_from_docx(file_path):
    """Extract all text from a .docx file."""
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_text_from_xlsx(file_path):
    """Extract all text from an .xlsx file, concatenating all sheets."""
    xls = pd.ExcelFile(file_path)
    chunks = []
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet, dtype=str)
        # Include column headers
        chunks.append(" ".join(str(c) for c in df.columns))
        for _, row in df.iterrows():
            chunks.append(" ".join(str(v) for v in row.values if pd.notna(v)))
    return "\n".join(chunks)


def parse_document(file_path):
    """Parse a .docx or .xlsx file and return raw text."""
    lower = file_path.lower()
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif lower.endswith(".xlsx"):
        return extract_text_from_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")


def tokenize_and_clean(text):
    """Lowercase, tokenize, remove stopwords, lemmatize."""
    tokens = re.findall(r"[a-z]{2,}", text.lower())
    return [
        LEMMATIZER.lemmatize(t)
        for t in tokens
        if t not in STOP_WORDS and len(t) > 2
    ]


def detect_themes(tokens):
    """Score each theme based on keyword overlap with tokens."""
    token_set = Counter(tokens)
    scores = {}
    for theme, seeds in THEME_SEEDS.items():
        score = sum(token_set.get(s, 0) for s in seeds)
        if score > 0:
            scores[theme] = score
    return scores


def extract_key_phrases(text, top_n=20):
    """Use TF-IDF on sentence-level chunks to find key phrases."""
    sentences = re.split(r"[.\n]+", text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    if len(sentences) < 2:
        return []
    vec = TfidfVectorizer(
        max_features=200, stop_words="english",
        ngram_range=(1, 3), min_df=1
    )
    tfidf = vec.fit_transform(sentences)
    scores = tfidf.sum(axis=0).A1
    vocab = vec.get_feature_names_out()
    ranked = sorted(zip(vocab, scores), key=lambda x: -x[1])
    return [phrase for phrase, _ in ranked[:top_n]]


def analyze_single_market(file_path, market_name):
    """Full analysis of one market's document."""
    raw = parse_document(file_path)
    tokens = tokenize_and_clean(raw)
    themes = detect_themes(tokens)
    phrases = extract_key_phrases(raw)
    return {
        "market": market_name,
        "raw_text": raw,
        "tokens": tokens,
        "themes": themes,
        "key_phrases": phrases,
        "word_count": len(tokens),
    }


def find_parallels(analyses):
    """
    Compare analyses from multiple markets and identify shared themes,
    overlapping phrases, and common patterns.
    """
    if len(analyses) < 2:
        return {"shared_themes": {}, "unique_themes": {}, "overlapping_phrases": [],
                "parallel_opportunities": []}

    # Shared themes: themes present in 2+ markets
    theme_presence = defaultdict(list)
    for a in analyses:
        for theme in a["themes"]:
            theme_presence[theme].append(a["market"])

    shared_themes = {
        theme: markets
        for theme, markets in theme_presence.items()
        if len(markets) >= 2
    }
    unique_themes = {
        theme: markets[0]
        for theme, markets in theme_presence.items()
        if len(markets) == 1
    }

    # Overlapping key phrases (present in 2+ markets)
    phrase_presence = defaultdict(list)
    for a in analyses:
        for phrase in a["key_phrases"]:
            phrase_presence[phrase].append(a["market"])

    overlapping_phrases = [
        {"phrase": phrase, "markets": markets}
        for phrase, markets in phrase_presence.items()
        if len(markets) >= 2
    ]

    # Build parallel opportunity descriptions
    parallel_opportunities = []
    for theme, markets in shared_themes.items():
        label = theme.replace("_", " ").title()
        market_list = ", ".join(m.title() for m in markets)
        # combine related phrases
        related = [
            p["phrase"] for p in overlapping_phrases
            if any(seed in p["phrase"] for seed in THEME_SEEDS.get(theme, []))
        ]
        parallel_opportunities.append({
            "theme": label,
            "markets": markets,
            "description": (
                f"{label} is a shared interest across {market_list}. "
                f"Related keywords: {', '.join(related[:5]) if related else 'general overlap'}."
            ),
        })

    return {
        "shared_themes": shared_themes,
        "unique_themes": unique_themes,
        "overlapping_phrases": overlapping_phrases,
        "parallel_opportunities": parallel_opportunities,
    }


# ---------------------------------------------------------------------------
# Profile & Campaign Generation
# ---------------------------------------------------------------------------

MARKET_CONTEXT = {
    "moldova": {
        "language": "Romanian / Russian",
        "platforms": ["Facebook", "Instagram", "Odnoklassniki", "TikTok"],
        "cultural_notes": (
            "Moldova has a young, increasingly tech-savvy population with strong "
            "ties to both Romanian and Russian-speaking internet spaces. "
            "Diaspora engagement is high; remittances drive consumer spending."
        ),
    },
    "georgia": {
        "language": "Georgian / Russian / English",
        "platforms": ["Facebook", "Instagram", "TikTok", "YouTube"],
        "cultural_notes": (
            "Georgia's tech scene is growing rapidly, with Tbilisi emerging as a "
            "digital nomad hub. Strong national identity and pride in heritage "
            "blend with openness to Western culture and EU aspirations."
        ),
    },
    "armenia": {
        "language": "Armenian / Russian / English",
        "platforms": ["Facebook", "Instagram", "Telegram", "TikTok"],
        "cultural_notes": (
            "Armenia has a vibrant startup ecosystem and a large, engaged diaspora. "
            "Tech education is highly valued. Community-driven campaigns resonate "
            "strongly; national pride and innovation go hand-in-hand."
        ),
    },
}


def build_audience_profile(parallel, market):
    """Create a target audience profile for a specific market based on a parallel."""
    ctx = MARKET_CONTEXT.get(market, {})
    theme = parallel["theme"]
    return {
        "market": market.title(),
        "theme": theme,
        "profile_name": f"{theme} Enthusiast — {market.title()}",
        "languages": ctx.get("language", "Local language"),
        "preferred_platforms": ctx.get("platforms", ["Facebook", "Instagram"]),
        "description": (
            f"Young professional or student in {market.title()} interested in "
            f"{theme.lower()}. Active on social media, consumes content in "
            f"{ctx.get('language', 'local language')}. Values practical skills, "
            f"career growth, and community connection."
        ),
        "cultural_context": ctx.get("cultural_notes", ""),
    }


def generate_campaign_ideas(parallel, market):
    """Generate social media campaign ideas for a given parallel + market."""
    ctx = MARKET_CONTEXT.get(market, {})
    theme = parallel["theme"]
    platforms = ctx.get("platforms", ["Facebook", "Instagram"])

    campaigns = []

    # Campaign 1: Educational content series
    campaigns.append({
        "name": f"{theme} Spotlight Series — {market.title()}",
        "platform": platforms[0] if platforms else "Facebook",
        "format": "Carousel posts / Short video series",
        "description": (
            f"A weekly content series highlighting {theme.lower()} trends, "
            f"success stories, and opportunities in {market.title()}. "
            f"Feature local voices and professionals to build credibility."
        ),
        "cta": f"Follow for weekly {theme.lower()} insights",
        "hashtags": [
            f"#{market.title()}{theme.replace(' ', '')}",
            f"#{theme.replace(' ', '')}",
            f"#{market.title()}Future",
        ],
    })

    # Campaign 2: Community challenge
    campaigns.append({
        "name": f"{theme} Challenge — {market.title()}",
        "platform": "TikTok" if "TikTok" in platforms else platforms[-1],
        "format": "User-generated content challenge",
        "description": (
            f"Launch a 7-day challenge encouraging users in {market.title()} "
            f"to share their {theme.lower()} journey. Offer prizes or "
            f"recognition for most engaging submissions."
        ),
        "cta": f"Join the #{market.title()}{theme.replace(' ', '')}Challenge",
        "hashtags": [
            f"#{market.title()}{theme.replace(' ', '')}Challenge",
            f"#{theme.replace(' ', '')}Goals",
        ],
    })

    # Campaign 3: Cross-market collaboration
    other_markets = [m for m in parallel["markets"] if m != market]
    if other_markets:
        partner = other_markets[0].title()
        campaigns.append({
            "name": f"{market.title()} × {partner}: {theme} Bridge",
            "platform": "Instagram" if "Instagram" in platforms else platforms[0],
            "format": "Collaborative live sessions / Joint posts",
            "description": (
                f"Cross-border collaboration between {market.title()} and "
                f"{partner} creators focused on {theme.lower()}. "
                f"Highlights shared challenges and aspirations to build "
                f"a regional community."
            ),
            "cta": f"Watch the {market.title()}×{partner} {theme.lower()} live",
            "hashtags": [
                f"#{market.title()}x{partner}",
                f"#Regional{theme.replace(' ', '')}",
            ],
        })

    return campaigns


def generate_all_outputs(analyses):
    """
    Main orchestrator: takes list of market analyses, finds parallels,
    and generates profiles + campaigns for each parallel × market combo.
    """
    parallels = find_parallels(analyses)
    results = {
        "parallels": parallels,
        "profiles": [],
        "campaigns": [],
    }

    for parallel in parallels["parallel_opportunities"]:
        for market in parallel["markets"]:
            profile = build_audience_profile(parallel, market)
            results["profiles"].append(profile)

            campaigns = generate_campaign_ideas(parallel, market)
            results["campaigns"].extend(campaigns)

    return results
