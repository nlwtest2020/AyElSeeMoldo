#!/usr/bin/env python3
"""
Stage 1: Parse a curriculum specification .docx file into structured JSON.

Usage:
    python scripts/parse_docx.py input/<filename>.docx

Output:
    output/parsed_curriculum.json

The parser extracts:
- Day/module title and duration
- Timing blocks with content topics
- SWBATs (Students Will Be Able To)
- Deliverables and assessments
- Notes and special instructions
"""

import json
import re
import sys
from pathlib import Path

from docx import Document


def extract_text_blocks(doc):
    """Extract all text from the document, preserving paragraph structure and styles."""
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        blocks.append({
            "text": text,
            "style": para.style.name if para.style else "Normal",
            "bold": any(run.bold for run in para.runs if run.bold),
            "level": para.style.name.split()[-1] if para.style and "Heading" in para.style.name else None,
        })
    return blocks


def extract_tables(doc):
    """Extract all tables from the document as lists of row dictionaries."""
    tables = []
    for table in doc.tables:
        rows = []
        headers = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = cells
                continue
            if headers:
                row_dict = {}
                for j, header in enumerate(headers):
                    if j < len(cells):
                        row_dict[header] = cells[j]
                rows.append(row_dict)
            else:
                rows.append({"cells": cells})
        tables.append({"headers": headers, "rows": rows})
    return tables


def parse_time_range(text):
    """Try to extract a time range from text like '9:00-10:30' or '9:00 AM - 10:30 AM'."""
    patterns = [
        r'(\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)?)\s*[-–]\s*(\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)?)',
        r'(\d{1,2}:\d{2})\s*[-–]\s*(\d{1,2}:\d{2})',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return {"start": match.group(1).strip(), "end": match.group(2).strip()}
    return None


def parse_duration(text):
    """Try to extract duration like '30 min', '1 hour', '1.5 hours'."""
    patterns = [
        r'(\d+(?:\.\d+)?)\s*(?:hours?|hrs?)',
        r'(\d+)\s*(?:minutes?|mins?)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0)
    return None


def find_swbats(blocks):
    """Identify SWBAT statements from text blocks."""
    swbats = []
    in_swbat_section = False
    swbat_patterns = [
        r'students?\s+will\s+be\s+able\s+to',
        r'SWBAT',
        r'swbat',
        r'learning\s+objectives?',
        r'objectives?:',
        r'by\s+the\s+end.*students?\s+will',
    ]

    for block in blocks:
        text = block["text"]
        # Check if this is the start of a SWBAT section
        for pattern in swbat_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                in_swbat_section = True
                # If the SWBAT is on the same line as the header
                remainder = re.split(pattern, text, flags=re.IGNORECASE)[-1].strip()
                if remainder and len(remainder) > 10:
                    swbats.append(remainder.lstrip(":- "))
                break

        if in_swbat_section and not any(re.search(p, text, re.IGNORECASE) for p in swbat_patterns):
            # Check if this looks like a SWBAT item (often bulleted or numbered)
            cleaned = text.lstrip("•-*0123456789.) ")
            if cleaned and len(cleaned) > 5:
                # Stop if we hit a new section header
                if block.get("level") or (block["bold"] and len(text) < 50):
                    in_swbat_section = False
                else:
                    swbats.append(cleaned)

    return swbats


def find_timing_blocks(blocks, tables):
    """Extract timing/schedule information from blocks and tables."""
    timing_blocks = []

    # Check tables first (schedules are often in tables)
    for table in tables:
        headers_lower = [h.lower() for h in table["headers"]]
        has_time = any("time" in h or "schedule" in h or "duration" in h for h in headers_lower)
        if has_time:
            for row in table["rows"]:
                block = {}
                for key, value in row.items():
                    key_lower = key.lower()
                    if "time" in key_lower or "schedule" in key_lower:
                        time_range = parse_time_range(value)
                        if time_range:
                            block["start_time"] = time_range["start"]
                            block["end_time"] = time_range["end"]
                        else:
                            block["time_raw"] = value
                        duration = parse_duration(value)
                        if duration:
                            block["duration"] = duration
                    elif "topic" in key_lower or "content" in key_lower or "activity" in key_lower or "section" in key_lower:
                        block["topic"] = value
                    elif "description" in key_lower or "detail" in key_lower or "notes" in key_lower:
                        block["description"] = value
                    elif "deliverable" in key_lower or "output" in key_lower:
                        block["deliverable"] = value
                    else:
                        block.setdefault("extra", {})[key] = value
                if block:
                    timing_blocks.append(block)

    # Also scan text blocks for time patterns
    if not timing_blocks:
        for block in blocks:
            text = block["text"]
            time_range = parse_time_range(text)
            if time_range:
                entry = {
                    "start_time": time_range["start"],
                    "end_time": time_range["end"],
                    "topic": text,
                }
                duration = parse_duration(text)
                if duration:
                    entry["duration"] = duration
                timing_blocks.append(entry)

    return timing_blocks


def find_deliverables(blocks):
    """Extract deliverables from text blocks."""
    deliverables = []
    in_section = False
    deliverable_patterns = [
        r'deliverables?',
        r'assignments?',
        r'outputs?',
        r'products?',
        r'submissions?',
    ]

    for block in blocks:
        text = block["text"]
        for pattern in deliverable_patterns:
            if re.search(pattern, text, re.IGNORECASE) and (block.get("level") or block["bold"]):
                in_section = True
                break

        if in_section:
            if block.get("level") and not any(re.search(p, text, re.IGNORECASE) for p in deliverable_patterns):
                in_section = False
            elif not block.get("level") and not block["bold"]:
                cleaned = text.lstrip("•-*0123456789.) ")
                if cleaned and len(cleaned) > 3:
                    deliverables.append(cleaned)

    return deliverables


def parse_docx(file_path):
    """Main parsing function. Returns structured curriculum data."""
    doc = Document(file_path)

    blocks = extract_text_blocks(doc)
    tables = extract_tables(doc)

    # Try to find the title (first heading or first bold text)
    title = "Untitled Curriculum Day"
    for block in blocks:
        if block.get("level") in ["1", "2"] or (block["bold"] and len(block["text"]) < 100):
            title = block["text"]
            break

    # Extract total duration from any mention of total time
    total_duration = None
    for block in blocks:
        duration = parse_duration(block["text"])
        if duration and ("total" in block["text"].lower() or "day" in block["text"].lower()):
            total_duration = duration
            break

    curriculum = {
        "title": title,
        "total_duration": total_duration,
        "swbats": find_swbats(blocks),
        "timing_blocks": find_timing_blocks(blocks, tables),
        "deliverables": find_deliverables(blocks),
        "tables": tables,
        "raw_sections": [],
    }

    # Group remaining content into sections by headings
    current_section = {"heading": "Introduction", "content": []}
    for block in blocks:
        if block.get("level"):
            if current_section["content"]:
                curriculum["raw_sections"].append(current_section)
            current_section = {"heading": block["text"], "content": []}
        else:
            current_section["content"].append(block["text"])
    if current_section["content"]:
        curriculum["raw_sections"].append(current_section)

    return curriculum


def main():
    if len(sys.argv) < 2:
        print("Usage: python scripts/parse_docx.py input/<filename>.docx")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    if not input_file.exists():
        print(f"Error: File not found: {input_file}")
        sys.exit(1)

    if not input_file.suffix.lower() == ".docx":
        print(f"Error: Expected .docx file, got {input_file.suffix}")
        sys.exit(1)

    script_dir = Path(__file__).parent.parent
    output_dir = script_dir / "output"
    output_dir.mkdir(exist_ok=True)
    output_file = output_dir / "parsed_curriculum.json"

    print(f"Parsing: {input_file}")
    curriculum = parse_docx(input_file)

    with open(output_file, "w") as f:
        json.dump(curriculum, f, indent=2)

    print(f"Output: {output_file}")
    print(f"  Title: {curriculum['title']}")
    print(f"  SWBATs found: {len(curriculum['swbats'])}")
    print(f"  Timing blocks: {len(curriculum['timing_blocks'])}")
    print(f"  Deliverables: {len(curriculum['deliverables'])}")
    print(f"  Sections: {len(curriculum['raw_sections'])}")

    if not curriculum["swbats"]:
        print("  WARNING: No SWBATs found. You may need to manually add them to the JSON.")
    if not curriculum["timing_blocks"]:
        print("  WARNING: No timing blocks found. You may need to manually add them to the JSON.")


if __name__ == "__main__":
    main()
