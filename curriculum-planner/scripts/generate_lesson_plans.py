#!/usr/bin/env python3
"""
Generate professional DOCX teaching guides for each session.

Produces scannable, EL Education-style lesson plans with clear teacher/student
actions, pacing notes, and embedded assessment checkpoints.

Usage:
    python scripts/generate_lesson_plans.py

Reads:  output/parsed_curriculum.json
        templates/ALC_Lesson_Plan_Template.docx
Output: output/Session_1_Lesson_Plan.docx (etc.)
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

SCRIPT_DIR = Path(__file__).parent.parent
OUTPUT_DIR = SCRIPT_DIR / "output"
TEMPLATE_PATH = SCRIPT_DIR / "templates" / "ALC_Lesson_Plan_Template.docx"
PARSED_PATH = OUTPUT_DIR / "parsed_curriculum.json"

# ALC Brand Colors
ALC_DARK_BLUE = RGBColor(0x24, 0x35, 0x54)
ALC_ACCENT_BLUE = RGBColor(0x2E, 0x74, 0xB5)
ALC_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ALC_LIGHT_GRAY = RGBColor(0x7F, 0x7F, 0x7F)
BREAK_BG = "D9E2F3"  # light blue-gray for breaks
HEADER_BG = "243554"  # dark blue for block headers
TIP_BG = "FFF2CC"  # warm yellow for teacher tips
CHECK_BG = "E2EFDA"  # soft green for CFU checks


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    existing = tcPr.findall(qn('w:shd'))
    for e in existing:
        tcPr.remove(e)
    shading = etree.SubElement(tcPr, qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell. Each border is a dict with sz, color, val."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for side, props in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if props:
            el = borders.find(qn(f'w:{side}'))
            if el is None:
                el = etree.SubElement(borders, qn(f'w:{side}'))
            el.set(qn('w:val'), props.get('val', 'single'))
            el.set(qn('w:sz'), str(props.get('sz', 4)))
            el.set(qn('w:color'), props.get('color', '000000'))
            el.set(qn('w:space'), '0')


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, qn('w:tblPr'))
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = borders.find(qn(f'w:{side}'))
        if el is None:
            el = etree.SubElement(borders, qn(f'w:{side}'))
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        el.set(qn('w:space'), '0')


def add_formatted_run(paragraph, text, bold=False, italic=False, size=None,
                      color=None, font_name=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run


def add_section_heading(doc, text, level=2):
    """Add a styled section heading."""
    p = doc.add_heading(text, level=level)
    if p.runs:
        p.runs[0].font.color.rgb = ALC_DARK_BLUE
    return p


def set_paragraph_spacing(paragraph, before=0, after=0):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


# ---------------------------------------------------------------------------
# GRR Phase Inference
# ---------------------------------------------------------------------------

def infer_grr_phase(format_text):
    """Infer Gradual Release of Responsibility phase from format description."""
    if not format_text:
        return ""
    ft = format_text.lower()

    phases = []
    # Check for I Do indicators
    if any(w in ft for w in ["instructor-led", "teacher-led", "lecture",
                              "interactive lecture", "demo", "direct instruction"]):
        phases.append("I Do")
    # Check for We Do indicators
    if any(w in ft for w in ["guided", "worked example", "full-group",
                              "full group", "debrief", "discussion"]):
        phases.append("We Do")
    # Check for You Do Together indicators
    if any(w in ft for w in ["pair", "pairs", "small group", "collaborative",
                              "group work", "partner"]):
        phases.append("You Do Together")
    # Check for You Do Alone indicators
    if any(w in ft for w in ["individual", "independent", "solo",
                              "self-assessment", "reflection"]):
        phases.append("You Do")

    if not phases:
        return ""
    return " \u2192 ".join(phases)


# ---------------------------------------------------------------------------
# Delivery Text Parsing
# ---------------------------------------------------------------------------

def parse_delivery_text(delivery):
    """Split delivery text into teacher actions and student actions."""
    if not delivery:
        return [], []

    # Split into sentences (handle periods followed by spaces, but not decimals)
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z"])', delivery)

    teacher_actions = []
    student_actions = []

    for sentence in sentences:
        s = sentence.strip()
        if not s:
            continue
        s_lower = s.lower()

        # Student action indicators
        if (s_lower.startswith("students ") or
            s_lower.startswith("each student") or
            s_lower.startswith("each pair") or
            s_lower.startswith("learners ") or
            (s_lower.startswith("they ") and student_actions)):
            student_actions.append(s)
        # Teacher action indicators or default
        else:
            teacher_actions.append(s)

    return teacher_actions, student_actions


def extract_quoted_phrases(text):
    """Extract quoted phrases from text (teacher scripting)."""
    return re.findall(r'\u201c([^\u201d]+)\u201d|"([^"]+)"', text)


# ---------------------------------------------------------------------------
# CFU Generation
# ---------------------------------------------------------------------------

def generate_cfu(activity_name, blooms, delivery):
    """Generate a check-for-understanding prompt based on the activity."""
    activity_lower = activity_name.lower()

    # Activity-specific checks
    if "prompt" in activity_lower or "craft" in activity_lower:
        return "Can students articulate why their revised prompt produced better output?"
    if "comparison" in activity_lower or "tool" in activity_lower:
        return "Can students name at least one specific difference between tool outputs?"
    if "how ai" in activity_lower or "works" in activity_lower:
        return "Can students explain in their own words what a token is and why it matters?"
    if "research" in activity_lower or "fact" in activity_lower:
        return "Are students checking sources, or accepting AI output at face value?"
    if "automat" in activity_lower or "workflow" in activity_lower:
        return "Can students identify the trigger, action, and output in their workflow?"
    if "capstone" in activity_lower or "portfolio" in activity_lower:
        return "Does the student's project solve a real problem using AI tools?"
    if "welcome" in activity_lower or "norms" in activity_lower:
        return "Are students engaged and clear on expectations for the day?"
    if "reflection" in activity_lower or "closing" in activity_lower:
        return "Can students name one specific thing they learned and one question they still have?"

    # Bloom's-level based fallback
    if blooms:
        blooms_lower = blooms.lower()
        if "create" in blooms_lower or "evaluate" in blooms_lower:
            return "Are student outputs showing original thinking, not just AI copy-paste?"
        if "analyze" in blooms_lower:
            return "Can students explain the reasoning behind their choices?"
        if "apply" in blooms_lower:
            return "Are students successfully completing the task independently?"

    return "Quick pulse check: thumbs up if confident, sideways if unsure, down if lost."


# ---------------------------------------------------------------------------
# Document Building
# ---------------------------------------------------------------------------

def generate_lesson_plan(session, template_doc):
    """Generate a professional teaching guide DOCX for one session."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Configure heading styles
    for level in range(1, 4):
        try:
            h_style = doc.styles[f"Heading {level}"]
            h_style.font.name = "Calibri"
            h_style.font.color.rgb = ALC_DARK_BLUE
        except KeyError:
            pass

    num = session["session_num"]
    title = session["title"]
    subtitle = session["subtitle"]
    schedule = session["schedule"]
    swbats = session["swbats"]

    # Parse subtitle
    day_info = ""
    time_info = ""
    theme_info = ""
    if subtitle:
        parts = subtitle.split("|")
        if len(parts) >= 1:
            day_info = parts[0].strip()
        if len(parts) >= 2:
            time_info = parts[1].strip()
        if len(parts) >= 3:
            theme_info = parts[2].strip().replace("Theme:", "").strip()

    instruction_blocks = [b for b in schedule if not b.get("is_break")]
    total_min = sum(int(b["duration"].replace(" min", "")) for b in instruction_blocks)

    # ======================================================================
    # SECTION 1: HEADER
    # ======================================================================
    heading = doc.add_heading(f"Session {num}: {title}", level=1)
    set_paragraph_spacing(heading, before=0, after=4)

    # Metadata table (borderless)
    meta_data = [
        ("Course", "AI & Digital Workflows Bootcamp"),
        ("Day / Time", f"{day_info}  |  {time_info}" if day_info else ""),
        ("Theme", theme_info),
        ("Duration", f"{total_min} min instruction  |  {len(instruction_blocks)} activity blocks"),
        ("Materials",
         "Slides deck, laptop/tablet per student, Claude & ChatGPT access, "
         "printed templates, whiteboard/markers"),
    ]

    meta_table = doc.add_table(rows=len(meta_data), cols=2)
    remove_table_borders(meta_table)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(meta_data):
        cell_label = meta_table.cell(i, 0)
        cell_value = meta_table.cell(i, 1)
        cell_label.text = ""
        cell_value.text = ""
        p = cell_label.paragraphs[0]
        add_formatted_run(p, label, bold=True, size=Pt(10), color=ALC_DARK_BLUE)
        p = cell_value.paragraphs[0]
        add_formatted_run(p, value, size=Pt(10))
        cell_label.width = Cm(3)
        cell_value.width = Cm(13)

    doc.add_paragraph()

    # ======================================================================
    # SECTION 2: LEARNING TARGETS
    # ======================================================================
    add_section_heading(doc, "LEARNING TARGETS")

    p = doc.add_paragraph()
    add_formatted_run(p, "By the end of this session, students will be able to:",
                      italic=True, size=Pt(10), color=ALC_LIGHT_GRAY)
    set_paragraph_spacing(p, after=4)

    for i, swbat in enumerate(swbats):
        # Clean up LO references
        clean = re.sub(r'\s*\(LO[^)]*\)', '', swbat).strip().rstrip(" ;,")
        p = doc.add_paragraph()
        add_formatted_run(p, f"{i + 1}. ", bold=True, size=Pt(10), color=ALC_ACCENT_BLUE)
        add_formatted_run(p, clean, size=Pt(10))
        set_paragraph_spacing(p, before=2, after=2)

    doc.add_paragraph()

    # ======================================================================
    # SECTION 3: ASSESSMENT AT A GLANCE
    # ======================================================================
    add_section_heading(doc, "ASSESSMENT AT A GLANCE")

    evidence = session.get("evidence", "")
    if evidence:
        p = doc.add_paragraph()
        add_formatted_run(p, "Look for: ", bold=True, size=Pt(10))
        add_formatted_run(p, evidence, size=Pt(10))
        set_paragraph_spacing(p, after=4)

    assess_items = [
        ("Formative",
         "Observe student work during each activity block. "
         "Check outputs match expectations in delivery notes."),
        ("Self-Assessment",
         "Students rate confidence (1\u20135) for each learning target at session close."),
        ("Exit Ticket",
         "See Closing section for specific questions."),
    ]
    for label, desc in assess_items:
        p = doc.add_paragraph()
        add_formatted_run(p, f"{label}: ", bold=True, size=Pt(10), color=ALC_DARK_BLUE)
        add_formatted_run(p, desc, size=Pt(10))
        set_paragraph_spacing(p, before=1, after=1)

    doc.add_paragraph()

    # ======================================================================
    # SECTION 4: LESSON FLOW
    # ======================================================================
    # Split into morning/afternoon
    morning_blocks = []
    afternoon_blocks = []
    lunch_seen = False

    for block in schedule:
        if "lunch" in block["activity"].lower():
            lunch_seen = True
            morning_blocks.append(block)
            continue
        if not lunch_seen:
            morning_blocks.append(block)
        else:
            afternoon_blocks.append(block)

    # Morning
    if morning_blocks:
        first_time = morning_blocks[0]["start_time"]
        last_time = morning_blocks[-1]["end_time"]
        add_section_heading(doc, f"MORNING BLOCK  ({first_time}\u2013{last_time})")
        for block in morning_blocks:
            add_activity_block(doc, block, session)

    # Afternoon
    if afternoon_blocks:
        doc.add_paragraph()
        first_time = afternoon_blocks[0]["start_time"]
        last_time = afternoon_blocks[-1]["end_time"]
        add_section_heading(doc, f"AFTERNOON BLOCK  ({first_time}\u2013{last_time})")
        for block in afternoon_blocks:
            add_activity_block(doc, block, session)

    doc.add_paragraph()

    # ======================================================================
    # SECTION 5: DIFFERENTIATION
    # ======================================================================
    add_section_heading(doc, "DIFFERENTIATION NOTES")

    diff_notes = get_differentiation_notes(num)
    for label, note in diff_notes:
        p = doc.add_paragraph()
        add_formatted_run(p, f"{label}: ", bold=True, size=Pt(10), color=ALC_DARK_BLUE)
        add_formatted_run(p, note, size=Pt(10))
        set_paragraph_spacing(p, before=2, after=2)

    doc.add_paragraph()

    # ======================================================================
    # SECTION 6: CLOSING & EXIT TICKET
    # ======================================================================
    add_section_heading(doc, "CLOSING & EXIT TICKET")

    exit_questions = get_exit_ticket_questions(num)
    p = doc.add_paragraph()
    add_formatted_run(p, "Ask students to write brief answers to:",
                      italic=True, size=Pt(10))
    set_paragraph_spacing(p, after=4)

    for i, question in enumerate(exit_questions):
        p = doc.add_paragraph()
        add_formatted_run(p, f"{i + 1}. ", bold=True, size=Pt(10))
        add_formatted_run(p, question, size=Pt(10))
        set_paragraph_spacing(p, before=2, after=2)

    doc.add_paragraph()

    # ======================================================================
    # SECTION 7: TEACHER REFLECTION
    # ======================================================================
    add_section_heading(doc, "POST-SESSION REFLECTION")

    prompts = [
        "What worked well today?",
        "What would I adjust for next time?",
        "Which students need follow-up or additional support?",
        "Did timing work? Where did I run long or short?",
    ]
    for prompt in prompts:
        p = doc.add_paragraph()
        add_formatted_run(p, f"\u25a2  {prompt}", size=Pt(10), color=ALC_LIGHT_GRAY)
        set_paragraph_spacing(p, before=2, after=6)
        # Add blank line for writing
        doc.add_paragraph()

    return doc


def add_activity_block(doc, block, session):
    """Render a single activity block as a visually distinct section."""
    is_break = block.get("is_break", False)
    activity = block["activity"]
    start = block["start_time"]
    end = block["end_time"]
    duration = block["duration"]

    # --- BREAK: simple divider ---
    if is_break:
        table = doc.add_table(rows=1, cols=1)
        remove_table_borders(table)
        cell = table.cell(0, 0)
        set_cell_shading(cell, BREAK_BG)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, f"{start}\u2013{end}  |  {activity}  |  {duration}",
                          bold=True, size=Pt(10), color=ALC_DARK_BLUE)
        set_paragraph_spacing(p, before=4, after=4)
        return

    format_str = block.get("format", "")
    blooms = block.get("blooms", "")
    delivery = block.get("delivery", "")
    pacing = block.get("pacing", "")
    instructor_note = block.get("instructor_note", "")
    grr = infer_grr_phase(format_str)

    # --- BLOCK HEADER (dark blue bar) ---
    header_table = doc.add_table(rows=1, cols=1)
    remove_table_borders(header_table)
    cell = header_table.cell(0, 0)
    set_cell_shading(cell, HEADER_BG)
    cell.text = ""

    # Line 1: Time | Activity | Duration
    p = cell.paragraphs[0]
    add_formatted_run(p, f"{start}\u2013{end}", bold=True, size=Pt(11), color=ALC_WHITE)
    add_formatted_run(p, f"    {activity}", bold=True, size=Pt(11), color=ALC_WHITE)
    add_formatted_run(p, f"    ({duration})", size=Pt(10), color=ALC_WHITE)

    # Line 2: Format | Bloom's | GRR
    meta_parts = []
    if format_str:
        # Shorten format for display
        fmt_short = format_str.split(".")[0].strip()
        if len(fmt_short) > 60:
            fmt_short = fmt_short[:57] + "..."
        meta_parts.append(fmt_short)
    if blooms:
        meta_parts.append(f"Bloom's: {blooms}")
    if grr:
        meta_parts.append(f"GRR: {grr}")

    if meta_parts:
        p = cell.add_paragraph()
        add_formatted_run(p, "  |  ".join(meta_parts), size=Pt(9), color=ALC_WHITE,
                          italic=True)
        set_paragraph_spacing(p, before=0, after=2)

    # --- BLOCK BODY (content table) ---
    body_table = doc.add_table(rows=1, cols=1)
    remove_table_borders(body_table)
    # Add left border for visual connection
    set_cell_borders(body_table.cell(0, 0),
                     left={'sz': 12, 'color': HEADER_BG, 'val': 'single'})
    body_cell = body_table.cell(0, 0)
    body_cell.text = ""

    # Parse delivery into teacher/student actions
    teacher_actions, student_actions = parse_delivery_text(delivery)

    # TEACHER DOES
    if teacher_actions:
        p = body_cell.paragraphs[0]
        add_formatted_run(p, "TEACHER DOES", bold=True, size=Pt(10),
                          color=ALC_ACCENT_BLUE)
        set_paragraph_spacing(p, before=6, after=2)

        for action in teacher_actions:
            p = body_cell.add_paragraph()
            # Check for quoted teacher scripting
            quotes = extract_quoted_phrases(action)
            if quotes:
                add_formatted_run(p, "\u2022  ", size=Pt(10))
                # Split around quotes to format them differently
                remaining = action
                for q_group in quotes:
                    q = q_group[0] or q_group[1]
                    parts = remaining.split(f'\u201c{q}\u201d', 1)
                    if len(parts) == 1:
                        parts = remaining.split(f'"{q}"', 1)
                    if len(parts) > 1:
                        add_formatted_run(p, parts[0], size=Pt(10))
                        add_formatted_run(p, f'\u201c{q}\u201d', size=Pt(10), italic=True)
                        remaining = parts[1] if len(parts) > 1 else ""
                    else:
                        add_formatted_run(p, remaining, size=Pt(10))
                        remaining = ""
                if remaining:
                    add_formatted_run(p, remaining, size=Pt(10))
            else:
                add_formatted_run(p, f"\u2022  {action}", size=Pt(10))
            set_paragraph_spacing(p, before=1, after=1)

    # STUDENTS DO
    if student_actions:
        p = body_cell.add_paragraph()
        add_formatted_run(p, "STUDENTS DO", bold=True, size=Pt(10),
                          color=ALC_ACCENT_BLUE)
        set_paragraph_spacing(p, before=6, after=2)

        for action in student_actions:
            p = body_cell.add_paragraph()
            add_formatted_run(p, f"\u2022  {action}", size=Pt(10))
            set_paragraph_spacing(p, before=1, after=1)

    # If no student actions were parsed but delivery exists, add a note
    if delivery and not student_actions and not teacher_actions:
        p = body_cell.paragraphs[0]
        add_formatted_run(p, delivery, size=Pt(10))

    # PACING
    if pacing:
        p = body_cell.add_paragraph()
        add_formatted_run(p, "\u23f1 PACING:  ", bold=True, size=Pt(9),
                          color=ALC_LIGHT_GRAY)
        add_formatted_run(p, pacing, size=Pt(9), color=ALC_LIGHT_GRAY)
        set_paragraph_spacing(p, before=6, after=2)

    # TEACHER TIP
    if instructor_note:
        # Use a mini-table for the tip background
        tip_table_in_cell = body_cell.add_paragraph()
        add_formatted_run(tip_table_in_cell, "\U0001f4a1 TEACHER TIP:  ", bold=True,
                          size=Pt(9), color=RGBColor(0x80, 0x60, 0x00))
        add_formatted_run(tip_table_in_cell, instructor_note, size=Pt(9), italic=True,
                          color=RGBColor(0x80, 0x60, 0x00))
        set_paragraph_spacing(tip_table_in_cell, before=4, after=2)

    # CHECK FOR UNDERSTANDING
    cfu = generate_cfu(activity, blooms, delivery)
    p = body_cell.add_paragraph()
    add_formatted_run(p, "\u2713 CHECK:  ", bold=True, size=Pt(9),
                      color=RGBColor(0x2D, 0x7D, 0x2D))
    add_formatted_run(p, cfu, size=Pt(9), color=RGBColor(0x2D, 0x7D, 0x2D),
                      italic=True)
    set_paragraph_spacing(p, before=4, after=6)

    # Small spacer after the block
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=2)


# ---------------------------------------------------------------------------
# Session-specific content
# ---------------------------------------------------------------------------

def get_differentiation_notes(session_num):
    """Return differentiation notes for each session."""
    return {
        1: [
            ("Struggling learners",
             "Provide a printed CRAFT template with fill-in-the-blank fields. "
             "Pair with a stronger partner during tool comparison. "
             "Offer a simplified prompt to start with."),
            ("Advanced learners",
             "Challenge them to chain multiple prompts together. "
             "Ask them to test edge cases (e.g., prompts that break the model). "
             "Have them document patterns they notice for the class."),
            ("Language support",
             "Allow prompts in students' first language, then translate. "
             "Provide key vocabulary list (token, temperature, prompt, output) "
             "with simple definitions."),
        ],
        2: [
            ("Struggling learners",
             "Provide a fact-checking checklist. Pre-select simpler research topics. "
             "Pair with a partner for source triangulation exercise."),
            ("Advanced learners",
             "Ask them to find a case where AI hallucination is subtle and hard to detect. "
             "Have them create their own trust rubric criteria."),
            ("Language support",
             "Allow use of translation tools alongside AI research tools. "
             "Provide sentence frames for analysis writing."),
        ],
        3: [
            ("Struggling learners",
             "Provide a pre-built automation template to modify rather than build from scratch. "
             "Use a step-by-step visual guide for workflow setup."),
            ("Advanced learners",
             "Challenge them to add error handling or branching logic. "
             "Ask them to document their workflow for a non-technical audience."),
            ("Language support",
             "Provide workflow vocabulary in context. "
             "Allow screen-recording of process instead of written documentation."),
        ],
        4: [
            ("Struggling learners",
             "Offer a capstone project template with sections pre-labeled. "
             "Allow a simpler scope (2 AI tools instead of 3+). "
             "Pair with a peer reviewer who gives encouraging, specific feedback."),
            ("Advanced learners",
             "Ask them to present their capstone process, not just the product. "
             "Challenge them to mentor a struggling peer. "
             "Have them write a 'Making Of' that could be published."),
            ("Language support",
             "Allow bilingual portfolio artifacts. "
             "Provide rubric in simplified language with examples."),
        ],
    }.get(session_num, [
        ("All learners", "Adjust pacing and grouping based on student needs."),
    ])


def get_exit_ticket_questions(session_num):
    """Return exit ticket questions tied to SWBATs for each session."""
    return {
        1: [
            "Name the 5 elements of the CRAFT framework and explain why one of them matters most to you.",
            "Give one specific example of when you'd choose Claude over ChatGPT (or vice versa), and why.",
            "What is one thing you learned today that changes how you'll use AI going forward?",
        ],
        2: [
            "Describe your process for fact-checking an AI-generated claim. What steps do you take?",
            "What is source triangulation, and why does using multiple AI tools matter?",
            "Rate your confidence (1-5) in distinguishing AI-generated text that needs human editing.",
        ],
        3: [
            "Draw or describe a simple automation workflow with a trigger, an action, and an output.",
            "What is one task in your daily work that could be automated? Sketch the workflow.",
            "Explain MCP (Model Context Protocol) in one sentence to a non-technical colleague.",
        ],
        4: [
            "What problem does your capstone project solve, and which AI tools did you use?",
            "What is one thing AI did well in your project, and one thing you had to fix or override?",
            "Name one specific, time-bound action you will take in the next 2 weeks to use AI at work.",
        ],
    }.get(session_num, [
        "What is the most important thing you learned today?",
        "What question do you still have?",
    ])


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    OUTPUT_DIR.mkdir(exist_ok=True)

    if not PARSED_PATH.exists():
        print(f"Error: {PARSED_PATH} not found. Run parse_docx.py first.")
        sys.exit(1)

    if not TEMPLATE_PATH.exists():
        print(f"Warning: Template not found: {TEMPLATE_PATH}. Using default styles.")

    with open(PARSED_PATH) as f:
        curriculum = json.load(f)

    # Load template for style reference (optional)
    template_doc = None
    if TEMPLATE_PATH.exists():
        template_doc = Document(TEMPLATE_PATH)

    for session in curriculum["sessions"]:
        num = session["session_num"]
        doc = generate_lesson_plan(session, template_doc)

        filename = f"Session_{num}_Lesson_Plan.docx"
        output_path = OUTPUT_DIR / filename
        doc.save(str(output_path))
        print(f"Generated: {output_path}")

    print(f"\nAll {len(curriculum['sessions'])} lesson plans generated.")


if __name__ == "__main__":
    main()
